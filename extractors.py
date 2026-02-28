"""
Document Ingestion & Text Extraction Module
============================================
Handles reading PDF and DOCX files, and provides a text-cleaning pipeline
to normalize the raw extracted text for downstream NLP processing.
"""

import re
import io
from pathlib import Path
from typing import List, Optional

import pdfplumber
from docx import Document


# ─── PDF Extraction (Layout Aware) ──────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using pdfplumber with dynamic column detection.
    Supports asymmetrical and full-width layouts.
    """
    extracted_text = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Extract words with their layout information
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False
            )

            if not words:
                continue

            # Identify if there's a strong vertical "gutter" (whitespace between columns)
            # and find the best split point.
            split_x = _find_column_split(words, page.width)

            if split_x:
                # Split words based on the dynamic gutter
                left_col = [w for w in words if w['x1'] <= split_x]
                right_col = [w for w in words if w['x0'] > split_x]
                
                # Heuristic: verify both columns have meaningful content
                if len(left_col) > 0.1 * len(words) and len(right_col) > 0.1 * len(words):
                    page_text = _format_column_words(left_col) + "\n" + _format_column_words(right_col)
                else:
                    page_text = _format_column_words(words)
            else:
                # Single column page
                page_text = _format_column_words(words)

            extracted_text.append(page_text)

    return "\n\n".join(extracted_text)


def _find_column_split(words: List[dict], page_width: float) -> Optional[float]:
    """
    Find the best vertical split point by looking for whitespace 'gutters'.
    Returns the X coordinate of the split, or None if no clear gutter exists.
    """
    # Create a density map of X-coordinates (1px resolution)
    # 0 = whitespace, >0 = has text
    x_density = [0] * int(page_width + 1)
    for w in words:
        start = max(0, int(w['x0']))
        end = min(int(page_width), int(w['x1']))
        for x in range(start, end + 1):
            x_density[x] += 1

    # Look for the widest contiguous zero-density (whitespace) lane
    # that is within the 20% to 80% range of the page width.
    best_gutter_start = 0
    best_gutter_width = 0
    
    current_start = -1
    search_start = int(page_width * 0.2)
    search_end = int(page_width * 0.8)

    for x in range(search_start, search_end):
        if x_density[x] == 0:
            if current_start == -1:
                # Ensure we have text to the left before starting a gutter
                if any(x_density[search_start:x]):
                    current_start = x
        else:
            if current_start != -1:
                width = x - current_start
                if width > best_gutter_width:
                    best_gutter_width = width
                    best_gutter_start = current_start
                current_start = -1
                
    # We don't check for gutter ending at search_end because it wouldn't be bounded on the right by text
    # within our search range.

    # If the gutter is significant (> 15px), return its center
    if best_gutter_width > 15:
        return best_gutter_start + (best_gutter_width / 2)
    
    return None


def _format_column_words(words: List[dict]) -> str:
    """Helper to group words into lines and format as text."""
    if not words:
        return ""

    # Sort words by top (Y) then x0 (X)
    words.sort(key=lambda x: (x['top'], x['x0']))

    lines = []
    current_line = [words[0]['text']]
    current_top = words[0]['top']

    for i in range(1, len(words)):
        word = words[i]
        # If the vertical difference is small, they are on the same line
        # We also check if the horizontal gap is reasonable
        if abs(word['top'] - current_top) < 3:
            current_line.append(word['text'])
        else:
            lines.append(" ".join(current_line))
            current_line = [word['text']]
            current_top = word['top']

    lines.append(" ".join(current_line))
    return "\n".join(lines)


# ─── DOCX Extraction ────────────────────────────────────────────────────────

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract raw text from a DOCX file using python-docx.
    Reads both paragraph text and table cell text.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text as a single string.
    """
    doc = Document(file_path)
    lines = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                lines.append(" | ".join(row_text))

    return "\n".join(lines)


# ─── Text Extraction Router ─────────────────────────────────────────────────

def extract_text(file_path: str) -> str:
    """
    Detect file type and extract text accordingly.

    Args:
        file_path: Path to the resume file (PDF, DOCX, or TXT).

    Returns:
        Cleaned text from the document.

    Raises:
        ValueError: If the file format is unsupported.
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        raw = extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        raw = extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
    else:
        raise ValueError(
            f"Unsupported file format: '{ext}'. Supported: .pdf, .docx, .txt"
        )

    cleaned = clean_text(raw)
    
    return cleaned


# ─── Text Cleaning Pipeline ─────────────────────────────────────────────────

def resolve_ligatures(text: str) -> str:
    """
    Resolve common PDF ligatures to standard ASCII characters.
    e.g. 'ﬁ' -> 'fi', 'ﬂ' -> 'fl'
    """
    ligatures = {
        "\ufb00": "ff",
        "\ufb01": "fi",
        "\ufb02": "fl",
        "\ufb03": "ffi",
        "\ufb04": "ffl",
        "\ufb05": "ft",
        "\ufb06": "st",
    }
    for k, v in ligatures.items():
        text = text.replace(k, v)
    return text


def strip_graphics_and_garbage(text: str) -> str:
    """
    Remove graphical artifacts, progress bars, and rating symbols.
    """
    if not text:
        return ""

    # 0. Resolve ligatures first
    text = resolve_ligatures(text)

    # 1. Remove repetitive progress bar characters (3 or more in a row)
    # e.g., "Python ------------------", "Java ..................."
    text = re.sub(r'[-_.=]{3,}', ' ', text)

    # 2. Strip common rating symbols (filled/empty circles, stars, etc.)
    # These often appear in skill sections as proficiency ratings.
    rating_symbols = [
        "\u25cf", "\u25cb", "\u2605", "\u2606", "\u25aa", "\u25ab",
        "\u25cf", "\u25cb", "\u2b24", "\u25ef", "\u25c6", "\u25c7",
        "\u2713", "\u2714", "\u25cf", "\u25ce", "\u25d0", "\u25d1"
    ]
    for symbol in rating_symbols:
        text = text.replace(symbol, " ")

    # 3. Aggressive garbage collection for non-standard Unicode
    # Replace anything that isn't a standard ASCII char, common punctuation, 
    # or whitespace with a space.
    text = re.sub(r'[^\x00-\x7F\s\u00A0-\u00FF]+', ' ', text)

    return text


def fix_spaced_text(text: str) -> str:
    """
    Repair text where characters are separated by single spaces.
    e.g., "J o h n   D o e" -> "John Doe"
    """
    # Look for sequences of single letters separated by spaces
    # We check for at least 3 characters in a row to avoid merging normal words
    return re.sub(r'(?<=\b\w) (?=\w\b)', '', text)


def clean_text(text: str) -> str:
    """
    Clean and normalize raw extracted text.
    """
    if not text:
        return ""

    # 1. Strip graphics and garbage symbols FIRST
    text = strip_graphics_and_garbage(text)

    # 2. Spacer fix for some PDF outputs (e.g., "J o h n")
    text = fix_spaced_text(text)

    # 2. Replace bullet characters with newlines
    bullet_chars = ["\u2022", "\u2023", "\u25aa", "\u25cf", "\u25cb",
                    "\u2013", "\u2014", "\u25e6", "\u2043", "\u00b7",
                    "\uf0b7", "\uf0a7"]
    for char in bullet_chars:
        text = text.replace(char, "\n")

    # 3. Replace smart quotes and dashes
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")

    # 4. Remove non-ASCII characters
    text = re.sub(r'[^\x00-\x7F\n]+', ' ', text)

    # 5. Normalize whitespace: collapse tabs and multiple spaces
    # BUT keep double spaces as they are used in sections
    text = re.sub(r'\t', '  ', text)
    text = re.sub(r' {3,}', '  ', text)

    # 6. Remove likely page numbers
    text = re.sub(r'^\s*\d{1,2}\s*$', '', text, flags=re.MULTILINE)

    # 7. Remove common header/footer noise
    text = re.sub(
        r'(?i)^(page\s*\d+.*|confidential.*|resume\s*of\s*.*|curriculum\s*vitae.*)$',
        '',
        text,
        flags=re.MULTILINE,
    )

    # 8. Collapse multiple blank lines into one
    text = re.sub(r'\n\s*\n', '\n\n', text)

    # 9. Strip leading/trailing whitespace
    text = text.strip()

    return text
